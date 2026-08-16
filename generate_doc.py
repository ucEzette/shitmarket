import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(f'''
            <w:tblBorders {nsdecls("w")}>
                <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:insideV w:val="none"/>
                <w:left w:val="none"/>
                <w:right w:val="none"/>
            </w:tblBorders>
        ''')
        tblPr[0].append(borders)

def build_doc():
    doc = docx.Document()
    
    # Page setup - Standard Letter, clean 0.65 in margins for a tight, crisp 1-2 page layout
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.header.is_linked_to_previous = False
    
    # Color definitions
    C_PRIMARY = RGBColor(15, 23, 42)      # Slate 900
    C_SECONDARY = RGBColor(37, 99, 235)   # Blue 600
    C_TEXT = RGBColor(51, 65, 85)         # Slate 700
    C_MUTED = RGBColor(100, 116, 139)     # Slate 500
    
    # Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(2)
    run_title = p_title.add_run("ShitMarket Protocol")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = C_PRIMARY
    
    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(12)
    run_sub = p_sub.add_run("Production-Grade Project Plan & Technical Delivery Specification")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = C_SECONDARY
    run_sub.font.bold = True
    
    # Divider line / Badges box
    p_badge = doc.add_paragraph()
    p_badge.paragraph_format.space_before = Pt(0)
    p_badge.paragraph_format.space_after = Pt(10)
    run_badge = p_badge.add_run("STATUS: Production Roadmap  |  TARGET CHAINS: Avalanche C-Chain / Solana  |  STANDARD: Institutional DeFi")
    run_badge.font.name = "Arial"
    run_badge.font.size = Pt(8.5)
    run_badge.font.color.rgb = C_MUTED
    run_badge.font.bold = True

    def add_section_header(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(title)
        run.font.name = "Arial"
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = C_PRIMARY
        return p

    # 1. Executive Summary
    add_section_header("1. Executive Summary & Value Proposition")
    p_exec = doc.add_paragraph()
    p_exec.paragraph_format.space_before = Pt(0)
    p_exec.paragraph_format.space_after = Pt(6)
    p_exec.paragraph_format.line_spacing = 1.15
    run_exec = p_exec.add_run(
        "ShitMarket is an institutional-grade, permissionless prediction market protocol built for high-velocity speculation on meme tokens, trending assets, and macro events. "
        "By synthesizing Two-Sided Constant Product Market Makers (CPMM) with hybrid CLOB order-routing and frictionless Web2-style social onboarding, ShitMarket delivers instant liquidity, sub-second execution, and tamper-proof multi-oracle settlement."
    )
    run_exec.font.name = "Arial"
    run_exec.font.size = Pt(9.5)
    run_exec.font.color.rgb = C_TEXT

    # Bullet Highlights
    bullets = [
        ("Instant Liquidity via CPMM:", " 1:1 USDC-backed ERC-1155 outcome tokens (MOON / JEET) with 0.10% swap fees (70% LP yield, 30% treasury)."),
        ("Zero-Friction UX:", " Gasless account abstraction paymasters, Privy embedded social logins, and universal multi-chain deposits via Circle CCTP."),
        ("Tamper-Proof Settlement:", " Multi-oracle TWAP aggregator (Chainlink, Pyth, Birdeye, DexScreener) paired with a 30-minute optimistic dispute & slashing window.")
    ]
    for bold_prefix, text in bullets:
        p_b = doc.add_paragraph(style='List Bullet')
        p_b.paragraph_format.space_before = Pt(0)
        p_b.paragraph_format.space_after = Pt(2)
        r_bold = p_b.add_run(bold_prefix)
        r_bold.font.name = "Arial"
        r_bold.font.size = Pt(9)
        r_bold.font.bold = True
        r_bold.font.color.rgb = C_PRIMARY
        r_txt = p_b.add_run(text)
        r_txt.font.name = "Arial"
        r_txt.font.size = Pt(9)
        r_txt.font.color.rgb = C_TEXT

    # 2. System Architecture
    add_section_header("2. Core Architecture & Technology Stack")
    
    table_arch = doc.add_table(rows=5, cols=3)
    table_arch.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_arch.autofit = False
    set_table_borders(table_arch, color="E2E8F0")

    col_widths = [Inches(1.5), Inches(1.8), Inches(3.8)]
    headers = ["Layer", "Technology", "Production Capabilities & Scope"]
    
    # Format Table Header
    for idx, name in enumerate(headers):
        cell = table_arch.cell(0, idx)
        cell.width = col_widths[idx]
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(name)
        run.font.name = "Arial"
        run.font.size = Pt(8.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    arch_data = [
        ("Smart Contracts", "Solidity 0.8.24 / Anchor (Rust)", "ShitMarketCore, AMPoolFactory (CPMM), ConditionalTokens (ERC-1155), DisputeBondEscrow."),
        ("Oracle & Dispute", "Pyth TWAP + UMA Optimistic Oracle", "Redundant median price aggregation, 20% outlier shield, on-chain bond challenge window."),
        ("Indexer & Services", "Node.js, TypeScript, Prisma, Redis", "Real-time event subscriptions, WebSocket streaming, automated keeper, cached leaderboards."),
        ("Frontend & Infra", "Next.js 14, Tailwind, Privy, AWS KMS", "Gasless transactions, Circle CCTP bridge, HSM KMS key custody, redundant RPC failover.")
    ]

    for row_idx, data in enumerate(arch_data, start=1):
        bg_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            cell = table_arch.cell(row_idx, col_idx)
            cell.width = col_widths[col_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(8.5)
            run.font.color.rgb = C_PRIMARY if col_idx == 0 else C_TEXT
            if col_idx == 0:
                run.font.bold = True

    # 3. Phased Execution Roadmap
    add_section_header("3. 8-Week Phased Execution Roadmap & Milestones")

    phases = [
        ("Phase 1: Vault Hardening & Security Audit (Weeks 1–2)", 
         "Execute 20,000+ mathematical fuzzing cycles on discrete payout curves and sequential claims. Complete third-party security audits (OtterSec/Sec3/OpenZeppelin standard) and configure Safe/Squads Multi-Sig administration."),
        ("Phase 2: Decentralized Oracle & Slashing Bonds (Weeks 3–4)", 
         "Deploy TWAP outlier-rejection engine across Pyth & Chainlink. Implement on-chain 30-minute challenge window with economic dispute bonds and arbitrator fee-slashing logic."),
        ("Phase 3: Key Custody & Cross-Chain Infrastructure (Weeks 5–6)", 
         "Migrate keeper private keys to AWS/GCP KMS with strict IAM rate-limits and automated gas refuel crons. Integrate Circle CCTP multi-chain liquidity ingestors (Ethereum, Base, Arbitrum, Solana)."),
        ("Phase 4: Mainnet Rollout & Scaling (Weeks 7–8)", 
         "Execute private Mainnet Alpha with capped TVL ($10k/room), launch public Immunefi Bug Bounty, and transition to permissionless production Mainnet GA.")
    ]

    for p_title_text, p_desc in phases:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r_title = p.add_run(p_title_text + "\n")
        r_title.font.name = "Arial"
        r_title.font.size = Pt(9)
        r_title.font.bold = True
        r_title.font.color.rgb = C_SECONDARY
        
        r_desc = p.add_run(p_desc)
        r_desc.font.name = "Arial"
        r_desc.font.size = Pt(8.5)
        r_desc.font.color.rgb = C_TEXT

    # 4. Risk Mitigation Matrix
    add_section_header("4. Risk Mitigation & Security Controls")

    table_risk = doc.add_table(rows=5, cols=4)
    table_risk.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_risk.autofit = False
    set_table_borders(table_risk, color="E2E8F0")

    risk_widths = [Inches(1.6), Inches(0.9), Inches(1.8), Inches(2.8)]
    risk_headers = ["Risk Vector", "Severity", "Potential Impact", "Production Mitigation Strategy"]

    for idx, name in enumerate(risk_headers):
        cell = table_risk.cell(0, idx)
        cell.width = risk_widths[idx]
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(name)
        run.font.name = "Arial"
        run.font.size = Pt(8.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    risk_data = [
        ("Oracle Manipulation", "CRITICAL", "False outcome resolution", "Multi-source TWAP averaging (5m/15m) + 20% outlier filter + 30m dispute challenge delay."),
        ("Sequential Claim Insolvency", "CRITICAL", "Late claimant payout failure", "Isolated escrow vaults per pool; discrete math with zero rounding deficit verified via fuzzing."),
        ("Keeper Key Compromise", "HIGH", "Unauthorized room settlement", "AWS KMS / HSM key storage with strict IP restrictions, IAM caller roles, and rate limiters."),
        ("Network RPC Congestion", "MEDIUM", "Failed / dropped transactions", "Dynamic priority fee escalators (ComputeBudget) + automated secondary RPC circuit-breaker.")
    ]

    for row_idx, data in enumerate(risk_data, start=1):
        bg_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            cell = table_risk.cell(row_idx, col_idx)
            cell.width = risk_widths[col_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(8)
            run.font.color.rgb = C_TEXT
            if col_idx == 0:
                run.font.bold = True
                run.font.color.rgb = C_PRIMARY
            elif col_idx == 1:
                run.font.bold = True
                run.font.color.rgb = RGBColor(220, 38, 38) if text == "CRITICAL" else RGBColor(217, 119, 6)

    # 5. Production Gate Criteria & KPIs
    add_section_header("5. Key Performance Indicators & Production Gate Criteria")
    
    kpis = [
        ("System Reliability:", " >= 99.95% API/WebSocket uptime; < 60s automated settlement latency post-expiration."),
        ("Economic Solvency:", " 0% bad debt; 100% full collateralization maintained across all active and settling pools."),
        ("Execution Performance:", " < 1.5s optimistic UI trade execution; sub-second indexer block-ingestion lag."),
        ("90-Day Growth Targets:", " $10M+ cumulative trading volume, 2,500+ permissionless rooms, 15,000+ active trader wallets.")
    ]

    for prefix, text in kpis:
        p_k = doc.add_paragraph(style='List Bullet')
        p_k.paragraph_format.space_before = Pt(0)
        p_k.paragraph_format.space_after = Pt(2)
        r_pre = p_k.add_run(prefix)
        r_pre.font.name = "Arial"
        r_pre.font.size = Pt(9)
        r_pre.font.bold = True
        r_pre.font.color.rgb = C_PRIMARY
        r_t = p_k.add_run(text)
        r_t.font.name = "Arial"
        r_t.font.size = Pt(9)
        r_t.font.color.rgb = C_TEXT

    # Save to disk
    output_path = "/Users/adam/Documents/shitmarket/ShitMarket_Protocol_Project_Plan.docx"
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

if __name__ == "__main__":
    build_doc()
